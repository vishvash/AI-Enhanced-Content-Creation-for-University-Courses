# -*- coding: utf-8 -*-
"""
Created on Fri Aug  9 01:38:45 2024

@author: Lenovo
"""

import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
# from llama_index.embeddings.gemini import GeminiEmbedding
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain.prompts import ChatPromptTemplate
from langchain.schema import StrOutputParser
from dotenv import load_dotenv
import re
import io
from docx import Document
from fpdf import FPDF
from docx.shared import Pt
import time
import random
import logging
from google.api_core.exceptions import ResourceExhausted

load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Configure logging
logging.basicConfig(filename='api_usage.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


def log_request(request_details):
    logging.info(f"Request: {request_details}")

def log_response(response_details):
    logging.info(f"Response: {response_details}")

def log_error(error_details):
    logging.error(f"Error: {error_details}")


def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")
    
prompt_template_qa = """
You are an expert question paper maker for technical fields. 
Generate new unique {batch_size} questions for which the answers should be present in the provided context. The questions should be of '{question_type}' based on the provided context and of plagiarism free.
After generating {batch_size} questions, start generating their respective answers.
Completion of each question/ answer should be mandatorily followed by the code "VISH"
Generate questions and answers as they are and Do NOT number, bullet, hyphen, or label any question or an answer.
If question type is "MCQ", questions should be generated with choices (A),(B),(C),(D) in a single line.
DO NOT generate answers for "Questions for Detailed Answers" question type.
Make sure Number of questions generated is equal to the number of respective answers generated, that means number of code VISH should be twice the number of batch_size. But this point is not applicable to "Questions for Detailed Answers"
Make sure all the Multiple-Choice questions are answered.

The format of the quiz could be one of the following:
- Multiple-choice:
    <Question>
    <Question>
    ....
    <Answer>
    <Answer>
    ....
    Example:
        What is the time complexity of a binary search tree? (A) O(n) (B) O(log n) (C) O(n^2) (D) O(1) VISH
        Which of the following is NOT a type of intelligent agent? (A) Simple reflex agent (B) Goal-based agent (C) Planning agent (D) Semi-autonomous agent VISH
        (B) O(log n) VISH
        (D) Semi-autonomous agent VISH
- True-false:
      <Question>
      <Question>
      .....
      <Answer>
      <Answer>
      .....
    Example:
        A binary search tree can only store unique values. VISH
        The root node of a binary search tree is always the smallest value in the tree. VISH
        True VISH
        False VISH
- Short Answer:
      <Question>
      <Question>
      ....
      <Answer>
      <Answer>
      ....
    Example:
        What is the time complexity of a binary search tree? VISH
        Describe the main advantage of using a binary search tree. VISH
        O(log n) VISH
        It provides efficient searching and sorting. VISH
- Detailed Answer:
      <Question>
      <Question>
      ....
      <Don't generate Answer> 
      <Don't generate Answer> 
      ....
    Example:
        Explain how binary search trees work and their applications. VISH
        Describe the process of balancing a binary search tree. VISH
       <No Answers> 
       <No Answers> 

   
context: {context}
    
"""

prompt_template_q = """
You are an expert question paper maker for technical fields. 
Generate new unique {batch_size} questions for which the answers should be present in the provided context. The questions should be of '{question_type}' based on the provided context and of plagiarism free.
Completion of each question should be mandatorily followed by the code "VISH"
Generate questions as they are and Do NOT number, bullet, hyphen, or label any question.
If question type is "Detailed Answer", don't generate answers for it.
If question type is "MCQ", questions should be generated with choices (A),(B),(C),(D) in a single line.

The format of the quiz could be one of the following:
- Multiple-choice:
    <Question>:
    <Question>:
    ....
    Example:
        What is the time complexity of a binary search tree? (A) O(n) (B) O(log n) (C) O(n^2) (D) O(1)
        Which of the following is NOT a type of intelligent agent? (A) Simple reflex agent (B) Goal-based agent (C) Planning agent (D) Semi-autonomous agent
- True-false:
      <Question>:
      <Question>:
      .....
    Example:
        A binary search tree can only store unique values.
        The root node of a binary search tree is always the smallest value in the tree.
- Short Answer:
      <Question>:
      <Question>:
      ....
    Example:
        What is the time complexity of a binary search tree?
        Describe the main advantage of using a binary search tree.
- Detailed Answer:
      <Question>:
      <Question>:
      ....
    Example:
        Explain how binary search trees work and their applications.
        Describe the process of balancing a binary search tree.

   
context: {context}
    
"""

def get_genquest_chain(prompt_template):
    prompt = ChatPromptTemplate.from_template(prompt_template)
    chain = prompt | ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3) | StrOutputParser()
    return chain

# Function to clean and trim questions
def clean_question(question):
    # Remove leading non-alphabetic characters and whitespaces
    return re.sub(r'^[^a-zA-Z]+', '', question).strip()

# Function to invoke the API with retry and exponential backoff
def invoke_with_retry(chain, request, max_retries=5, initial_delay=1):
    delay = initial_delay
    for attempt in range(max_retries):
        try:
            # Log the request details
            log_request(request)
            # Attempt the API call
            response = chain.invoke(request)
            # Log the response details
            log_response(response)
            return response
        except ResourceExhausted as e:
            log_error(f"Attempt {attempt + 1} failed with ResourceExhausted error: {e}")
            if attempt < max_retries - 1:
                # If not the last attempt, wait before retrying
                time.sleep(delay)
                delay *= 2  # Exponential backoff
            else:
                # If all attempts fail, raise the exception
                log_error(f"Max retries reached. Raising exception: {e}")
                raise e


# Function to generate questions iteratively
def generate_questions(qa, question_type, num_questions, topic, batch_size=5):
    questions = []
    answers = []
    template = prompt_template_qa if qa == 1 else prompt_template_q

    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings)
    docs = new_db.similarity_search(topic)

    chain = get_genquest_chain(template)
        
    question_type = {"MCQ": "MCQs along with choices (A),(B),(C),(D) in single line", "True/False": "Questions for True/False QA", "Short Answer": "Questions for Short QA", "Detailed Answer": "Questions for Detailed QA"}.get(question_type)

    while len(questions) < num_questions:
        # Prepare the request
        request = {
            "batch_size": batch_size,
            "question_type": question_type,
            "context": docs
        }

        # Call the API with retry logic
        response = invoke_with_retry(chain, request)
       
        
        # Split response by VISH to separate questions and answers sections
        parts = response.split("VISH")
        
        if qa == 1:
            
            answer_start_index = batch_size
    
            # Separate questions and answers
            generated_questions = parts[:answer_start_index]
            generated_answers = parts[answer_start_index:]
    
            # Clean questions and answers
            generated_questions = [clean_question(q) for q in generated_questions if q.strip()]
            generated_answers = [a.strip() for a in generated_answers if a.strip()]
            
            questions.extend(generated_questions)
            answers.extend(generated_answers)

        else:
            generated_questions = [clean_question(q) for q in parts if q.strip()]
            questions.extend(generated_questions)        

        # Delay to avoid crossing API rate limits
        time.sleep(random.uniform(1, 3))  # Random delay between 1 to 3 seconds
        

    # Truncate lists to requested number of questions
    numbered_questions = [f"{i + 1}. {q}" for i, q in enumerate(questions[:num_questions])]

    generated_answers = answers[:num_questions] if qa == 1 and question_type != "Detailed Answer" else []
    
    return numbered_questions, generated_answers


def get_conversational_chain():
    prompt_template = """
    Generate the answer as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
    provided context just say, "answer is not available in the context", don't provide the wrong answer\n\n
    Context:\n {context}?\n
    Question: \n{question}\n
    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)


    return chain

def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings)
    docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain()
    response = chain({"input_documents": docs, "question": f"Explain in detail: {user_question}"}, return_only_outputs=True)
    return response["output_text"]

def generate_docx(questions_by_type, answers_by_type, qa):
    doc = Document()
    doc.add_heading("Generated Questions and Answers", level=1)
    for question_type, questions in questions_by_type.items():
        doc.add_heading(f"\n{question_type}", level=1)
        if question_type == "Detailed Answer":
            for question in questions:
                doc.add_paragraph(f"{question}")
                if qa == 1:
                    answer = st.session_state['detailed_answers'].get(question, "")
                    if answer:
                        para = doc.add_paragraph()
                        para.add_run(f"Answer: {answer}")
                        para.paragraph_format.left_indent = Pt(24)  # 24 points is equivalent to 4 spaces
        else:
            for i, question in enumerate(questions):
                doc.add_paragraph(f"{question}")
                if qa == 1:
                    answer = answers_by_type.get(question_type, [])[i] if i < len(answers_by_type.get(question_type, [])) else "Answer not available"
                    doc.add_paragraph(f"   Answer: {answer}")
    return doc

# Function to replace unsupported characters
def replace_unsupported_characters(text):
    replacements = {
        '\u2013': '-',  # En dash replaced with hyphen
        '\u2022': '-',  # Bullet point replaced with hyphen
        # Add other replacements if needed
    }
    for key, value in replacements.items():
        text = text.replace(key, value)
    return text

def generate_pdf(questions_by_type, answers_by_type, qa):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, txt="Generated Questions and Answers", ln=True, align='L')
    pdf.ln(10)  # Add some space before the first heading
    
    for question_type, questions in questions_by_type.items():
        pdf.set_font("Arial", 'B', 14)  # Set a bold font for headings
        pdf.cell(0, 10, txt=question_type, ln=True, align='L')
        pdf.ln(5)  # Add a small space after the heading
        
        pdf.set_font("Arial", size=12)
        if question_type == "Detailed Answer":
            for question in questions:
                pdf.cell(0, 10, txt=replace_unsupported_characters(question), ln=True, align='L')
                if qa == 1:
                    answer = st.session_state['detailed_answers'].get(question, "")
                    if answer:
                        pdf.set_x(pdf.get_x() + 10)  # Indent by 10 units for the answer
                        pdf.multi_cell(0, 10, txt=f"Answer: {replace_unsupported_characters(answer)}")
                        pdf.set_x(pdf.r_margin)
                    pdf.ln(5)
        else:
            for i, question in enumerate(questions):
                pdf.multi_cell(0, 10, txt=f"{replace_unsupported_characters(question)}")
                if qa == 1:
                    answer = answers_by_type.get(question_type, [])[i] if i < len(answers_by_type.get(question_type, [])) else "Answer not available"
                    pdf.multi_cell(0, 10, txt=f"   Answer: {replace_unsupported_characters(answer)}")
                pdf.ln(5)  # Add space between Q&A
    return pdf

def convert_to_bytes(file_obj):
    buf = io.BytesIO()
    file_obj.save(buf)
    buf.seek(0)
    return buf.read()


def main():
    # Set the page config with a catchy title and favicon (optional)
    st.set_page_config(page_title="Gemini: Smart PDF Q&A Generator & Interactive Chatbox", page_icon="📝")
    
    image_url = r"https://i.pinimg.com/originals/29/56/51/295651e0ebc2d447aa9752a9a2fd9aa7.jpg"
    
    # Set a background image with a semi-transparent overlay
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-image: url('{image_url}');
            background-size: cover;
            background-repeat: no-repeat;
            background-attachment: fixed;
            background-position: center;
            color: #FFFFFF;
            font-family: 'Arial', sans-serif;
        }}

        </style>
        <div class="overlay"></div>
        """,
        unsafe_allow_html=True
    )

    st.header("Gemini: Smart PDF Q&A Generator & Interactive Chatbox")

    # Initialize or retrieve session state variables
    if 'raw_text' not in st.session_state:
        st.session_state['raw_text'] = ""
    if 'generated_questions' not in st.session_state:
        st.session_state['generated_questions'] = {}  # Initialize as an empty dictionary
    if 'generated_answers' not in st.session_state:
        st.session_state['generated_answers'] = {}  # Initialize as an empty dictionary
    if 'detailed_answers' not in st.session_state:
        st.session_state['detailed_answers'] = {}
    if 'qa' not in st.session_state:
        st.session_state['qa'] = None

    # Sidebar menu
    with st.sidebar:
        st.title("Menu:")
        st.markdown("## 📚 Upload your PDF Files and Click on the Submit & Process Button")
        pdf_docs = st.file_uploader("Upload PDF Files", accept_multiple_files=True)
        

        
        if st.button("Submit & Process"):
            if pdf_docs:
                with st.spinner("Processing..."):                   
                    st.session_state['raw_text'] = get_pdf_text(pdf_docs)
                    text_chunks = get_text_chunks(st.session_state['raw_text'])
                    get_vector_store(text_chunks)
                    st.success("Processing completed!")
            else:
                st.warning("Please upload at least one PDF file.")

    # Select question types
    st.subheader("⚙️ Generate Questions")

    # Create columns for the buttons
    col1_qa, col2_q = st.columns(2)
    
    with col1_qa:
        with_answers = st.session_state.get('qa') == 1
        if st.button("With Answers", key="with_answers"):
            st.session_state['qa'] = 1
    
        # Display text to show which option is selected
        if with_answers:
            st.markdown("**Get Comprehensive Q&As**")
    
    with col2_q:
        without_answers = st.session_state.get('qa') == 0
        if st.button("Without Answers", key="without_answers"):
            st.session_state['qa'] = 0
    
        # Display text to show which option is selected
        if without_answers:
            st.markdown("**Get Quick Questions**")

        
    question_types = ["MCQ", "True/False", "Short Answer", "Detailed Answer"]
    selected_question_types = st.multiselect("Select Question Types", question_types)
    num_questions = st.selectbox("Select Number of Questions", list(range(1, 101)))
    topic_key = "topic"
    topic = st.text_input("Enter the topic", key=topic_key)

    # Generate questions based on input
    if st.button("Generate Questions"):
        if st.session_state['raw_text']:
            if st.session_state['qa'] is not None:
                with st.spinner("Generating..."):
                    # context = st.session_state['raw_text']  # Use the processed PDF text as the context
                    qa = st.session_state['qa']
                    for question_type in selected_question_types:
                        questions, answers = generate_questions(qa, question_type, num_questions, topic)
                        st.session_state['generated_questions'][question_type] = questions
                        st.session_state['generated_answers'][question_type] = answers
                st.success("Questions generated successfully!") 
            else:
                st.warning("Please select either 'With Answers' or 'Without Answers'.")
        else:
            st.warning("Please process the PDF files first.")

    # Display generated questions and answers
    for question_type in selected_question_types:
        if question_type in st.session_state['generated_questions']:
            st.write(f"## {question_type}\n")
            
            questions = st.session_state['generated_questions'][question_type]
            
            if st.session_state['qa'] == 1:
                answers = st.session_state['generated_answers'].get(question_type, [])
                
                for i, question in enumerate(questions):                   
                    explain_key = f"explain_{i}_{question_type}"
                    
                    if question_type == "Detailed Answer":
                        # Create a unique key for each expander
                        expander_key = f"expander_{i}_{question_type}"
                    
                        # Check if this expander is open
                        if expander_key not in st.session_state:
                            st.session_state[expander_key] = False
                        
                        # Use the expander's key to keep it open or closed based on user interaction
                        with st.expander(f"**{question}**", expanded=st.session_state[expander_key]):
                            # Check if the answer already exists in session state
                            detailed_answer = st.session_state['detailed_answers'].get(question)
                            
                            if detailed_answer:
                                st.markdown(f"{detailed_answer}")
                            
                            if st.button("Explain", key=explain_key):
                                # Generate a new detailed answer and update session state
                                detailed_answer = user_input(question)
                                st.session_state['detailed_answers'][question] = detailed_answer
                                st.markdown(f"{detailed_answer}")
                                
                                # Set the expander to open
                                st.session_state[expander_key] = True
    
                            
                    else:
                        if i < len(answers):
                            with st.expander(f"**{question}**"):
                                st.markdown(f"{answers[i]}")
                        else:
                            st.warning("Answer not available for this question.")
            else:
                for question in questions:
                    st.markdown(question)

    # Download options
    st.subheader("📥 Download Options")
    for question_type in selected_question_types:
        if question_type in st.session_state['generated_questions']:
            questions = st.session_state['generated_questions'][question_type]
            answers = st.session_state['generated_answers'].get(question_type, [])
    
            # Define columns for main button and download options
            col1, col2 = st.columns([3, 1])  # Wider column for main button, narrower for download options
            
            with col1:
                if st.button(f"Download {question_type} Questions and Answers", key=f"download_{question_type}"):
                    doc = generate_docx({question_type: questions}, {question_type: answers}, st.session_state['qa'])
                    doc_bytes = convert_to_bytes(doc)
                    pdf = generate_pdf({question_type: questions}, {question_type: answers}, st.session_state['qa'])
                    pdf_bytes = pdf.output(dest='S').encode('latin1')
    
                    # Define inner columns within col2 for side-by-side buttons
                    with col2:
                        col1_dl, col2_dl = st.columns(2)  # Equal width for download buttons
                        with col1_dl:
                            st.download_button(
                                label="Docx",
                                data=doc_bytes,
                                file_name=f"{question_type.lower().replace(' ', '_')}_questions_answers.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        with col2_dl:
                            st.download_button(
                                label="PDF",
                                data=pdf_bytes,
                                file_name=f"{question_type.lower().replace(' ', '_')}_questions_answers.pdf",
                                mime="application/pdf"
                            )
    # Define columns for main button and download options
    col1, col2 = st.columns([3, 1])  # Wider column for main button, narrower for download options
    
    with col1:
        
        # Option to download all selected question types
        if st.button("Download All Q&A", key="download_all"):
            all_questions_by_type = {}
            all_answers_by_type = {}
        
            for question_type in selected_question_types:
                questions = st.session_state['generated_questions'].get(question_type, [])
                answers = st.session_state['generated_answers'].get(question_type, [])
                all_questions_by_type[question_type] = questions
                all_answers_by_type[question_type] = answers
        
            if all_questions_by_type:
                doc = generate_docx(all_questions_by_type, all_answers_by_type, st.session_state['qa'])
                doc_bytes = convert_to_bytes(doc)
                pdf = generate_pdf(all_questions_by_type, all_answers_by_type, st.session_state['qa'])
                pdf_bytes = pdf.output(dest='S').encode('latin1')

                # Define inner columns within col2 for side-by-side buttons
                with col2:
                    col1_dl, col2_dl = st.columns(2)  # Equal width for download buttons
                    with col1_dl:
                        st.download_button(
                            label="Docx",
                            data=doc_bytes,
                            file_name="all_selected_questions_answers.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    with col2_dl:
                        st.download_button(
                            label="PDF",
                            data=pdf_bytes,
                            file_name="all_selected_questions_answers.pdf",
                            mime="application/pdf"
                        )
            else:
                st.warning("No questions available for download.")
    
    
    # Chatbot input
    st.subheader("💬 Ask a Question")
    user_question_key = "user_question"
    user_question = st.text_input("Ask a Question from the PDF Files", key=user_question_key)
    if user_question:
        st.write(user_input(user_question))

if __name__ == "__main__":
    main()
